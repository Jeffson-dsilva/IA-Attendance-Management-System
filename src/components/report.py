from docx import Document

# Create a new Word document
doc = Document()

# Title Page
doc.add_heading("Internal Assessment & Attendance Management System", level=1)
doc.add_paragraph("\nMini Project Report\n\nSubmitted in partial fulfillment for the requirement of\nMaster of Computer Applications (MCA)\n\nDepartment of Computer Applications")
doc.add_paragraph("\nSubmitted by: [Your Name]\nRegister Number: [Your Register Number]")
doc.add_paragraph("\nUnder the guidance of: [Guide Name]\n\n[Institution Name]\n[Year]")

# 1. Synopsis
doc.add_heading("1. Synopsis", level=1)
doc.add_heading("1.1 Project Overview", level=2)
doc.add_paragraph("The Internal Assessment & Attendance Management System is a web-based platform designed to streamline academic processes for students and faculty. It allows efficient management of attendance records, internal assessment (IA) marks, and report generation while ensuring role-based access for different users.")

doc.add_heading("1.2 Scope", level=2)
doc.add_paragraph("The system includes functionalities such as student registration, faculty assignment, attendance tracking, IA marks management, and reporting. It does not cover external examinations or administrative activities outside academic evaluations.")

doc.add_heading("1.3 Target Audience", level=2)
doc.add_paragraph("This system is designed for students, faculty members, Head of Department (HOD), and examination authorities in educational institutions.")

doc.add_heading("1.4 Technologies Used", level=2)
doc.add_paragraph("Frontend: React.js, Tailwind CSS\nBackend: Node.js with Express.js\nDatabase: MongoDB\nAdditional Tools: Chart.js for analytics, Recharts for visualization")

# 2. Requirements
doc.add_heading("2. Requirements", level=1)
doc.add_heading("2.1 Functional Requirements", level=2)
doc.add_paragraph("• User authentication (students, faculty, HOD, admin)\n• Attendance tracking with real-time updates\n• IA Marks entry and modification by faculty\n• Automated report generation for attendance and marks\n• Role-based dashboard for different users")

doc.add_heading("2.2 Non-Functional Requirements", level=2)
doc.add_paragraph("• Security: Encrypted user authentication\n• Scalability: Supports multiple institutions\n• Usability: User-friendly interface\n• Performance: Optimized database queries")

doc.add_heading("2.3 System Requirements", level=2)
doc.add_paragraph("• Operating System: Windows/Linux/MacOS\n• Browser: Chrome, Firefox, Edge\n• Server: Node.js runtime\n• Database: MongoDB")

doc.add_heading("2.4 Third-Party Tools/Libraries", level=2)
doc.add_paragraph("• bcrypt.js for password hashing\n• JWT for authentication\n• Multer for file uploads\n• Nodemon for development server monitoring")

# 3. System Design
doc.add_heading("3. System Design", level=1)
doc.add_heading("3.1 Database Design", level=2)
doc.add_paragraph("The system uses MongoDB with collections such as 'Users', 'Attendance', and 'IAMarks'. User roles determine access to specific collections.")

doc.add_heading("3.2 User Interface Design", level=2)
doc.add_paragraph("The UI is built with React.js and Tailwind CSS, providing a responsive layout for different user roles. Dashboards include attendance charts, IA marks visualization, and real-time updates.")

# 4. Implementation
doc.add_heading("4. Implementation", level=1)
doc.add_heading("4.1 Detailed Steps", level=2)
doc.add_paragraph("• Step 1: Set up the backend with Node.js and Express.js\n• Step 2: Design MongoDB schema for users, attendance, and marks\n• Step 3: Develop the React.js frontend with API integration\n• Step 4: Implement authentication and role-based access control\n• Step 5: Integrate data visualization with Recharts\n• Step 6: Conduct testing and optimization")

doc.add_heading("4.2 Code Snippets", level=2)
doc.add_paragraph("Example API for fetching attendance data:\n\n```javascript\nrouter.get('/student-attendance/:email', async (req, res) => {\n  const { email } = req.params;\n  const student = await User.findOne({ email });\n  const attendanceRecords = await Attendance.find({ usn: student.usn });\n  res.json({ success: true, attendanceRecords });\n});\n```")

doc.add_heading("4.3 Testing", level=2)
doc.add_paragraph("• Unit Testing: Jest for backend API testing\n• Integration Testing: Postman for API validation\n• User Acceptance Testing: Conducted with faculty and students")

# 5. Results and Discussion
doc.add_heading("5. Results and Discussion", level=1)
doc.add_heading("5.1 Functionality", level=2)
doc.add_paragraph("The system successfully allows role-based access and provides real-time updates for attendance and IA marks.")

doc.add_heading("5.2 Usability", level=2)
doc.add_paragraph("The UI is intuitive, with simple navigation for students and faculty.")

doc.add_heading("5.3 Limitations", level=2)
doc.add_paragraph("• Requires internet connectivity\n• Faculty must manually enter IA marks")

doc.add_heading("5.4 Future Work", level=2)
doc.add_paragraph("• Mobile application integration\n• AI-based student performance prediction\n• Automated attendance tracking using face recognition")

# 6. Screenshots (Placeholder Section)
doc.add_heading("6. Screenshots", level=1)
doc.add_paragraph("Add screenshots of login page, student dashboard, faculty dashboard, IA marks entry, and attendance graphs here.")

# 7. References
doc.add_heading("7. References", level=1)
doc.add_paragraph("• React.js Documentation: https://react.dev\n• MongoDB Documentation: https://www.mongodb.com/docs\n• Node.js Express Guide: https://expressjs.com/")

# Save the document
file_path = "/mnt/data/Internal_Assessment_Attendance_Report.docx"
doc.save(file_path)

# Provide the download link
file_path
